#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Media ATLAS AI Telegram Bot - Single User with Media Features
Voice, Documents, Notes - No Multi-User Complexity
"""

import os
import sys
import time
import requests
from datetime import datetime
import tempfile
import subprocess

class MediaAtlasBot:
    """Media ATLAS AI Bot - Single User with Media Capabilities"""
    
    def __init__(self):
        self.bot_token = os.getenv('TELEGRAM_BOT_TOKEN')
        self.groq_api_key = os.getenv('GroqAPIKey')
        self.assistant_name = os.getenv('AssistantName', 'ATLAS')
        self.creator_name = os.getenv('Creator', 'K.V.SARVESH')
        self.last_update_id = 0
        
        if not self.bot_token or not self.groq_api_key:
            print("❌ Missing required environment variables")
            sys.exit(1)
    
    def send_message(self, chat_id: int, text: str) -> bool:
        """Send message via Telegram API"""
        try:
            response = requests.post(
                f"https://api.telegram.org/bot{self.bot_token}/sendMessage",
                json={
                    'chat_id': chat_id,
                    'text': text,
                    'parse_mode': 'HTML'
                },
                timeout=30
            )
            return response.status_code == 200
        except:
            return False
    
    def send_voice_message(self, chat_id: int, voice_file_path: str) -> bool:
        """Send voice message via Telegram API"""
        try:
            with open(voice_file_path, 'rb') as voice_file:
                response = requests.post(
                    f"https://api.telegram.org/bot{self.bot_token}/sendVoice",
                    data={
                        'chat_id': chat_id
                    },
                    files={
                        'voice': voice_file
                    },
                    timeout=30
                )
            return response.status_code == 200
        except Exception as e:
            print(f"❌ Voice send error: {e}")
            return False
    
    def send_document(self, chat_id: int, document_path: str, caption: str = "") -> bool:
        """Send document via Telegram API"""
        try:
            with open(document_path, 'rb') as doc_file:
                response = requests.post(
                    f"https://api.telegram.org/bot{self.bot_token}/sendDocument",
                    data={
                        'chat_id': chat_id,
                        'caption': caption
                    },
                    files={
                        'document': doc_file
                    },
                    timeout=30
                )
            return response.status_code == 200
        except Exception as e:
            print(f"❌ Document send error: {e}")
            return False
    
    def text_to_speech(self, text: str) -> str:
        """Convert text to speech using pyttsx3"""
        try:
            import pyttsx3
            engine = pyttsx3.init()
            
            # Create temporary file
            temp_file = tempfile.NamedTemporaryFile(suffix='.mp3', delete=False)
            temp_file.close()
            
            # Save speech to file
            engine.save_to_file(text, temp_file.name)
            engine.runAndWait()
            
            return temp_file.name
        except Exception as e:
            print(f"❌ TTS error: {e}")
            return None
    
    def generate_pdf_document(self, title: str, content: str) -> str:
        """Generate PDF document"""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            
            # Create temporary file
            temp_file = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
            temp_file.close()
            
            # Create PDF
            c = canvas.Canvas(temp_file.name, pagesize=letter)
            c.setFont("Helvetica-Bold", 16)
            c.drawString(50, 750, title)
            c.setFont("Helvetica", 12)
            
            # Add content (simple text wrapping)
            y_position = 700
            lines = content.split('\n')
            for line in lines:
                if y_position < 50:
                    c.showPage()
                    y_position = 750
                c.drawString(50, y_position, line)
                y_position -= 20
            
            c.save()
            return temp_file.name
        except Exception as e:
            print(f"❌ PDF error: {e}")
            return None
    
    def generate_word_document(self, title: str, content: str) -> str:
        """Generate Word document"""
        try:
            from docx import Document
            
            # Create temporary file
            temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
            temp_file.close()
            
            # Create Word document
            doc = Document()
            doc.add_heading(title, 0)
            doc.add_paragraph(content)
            doc.save(temp_file.name)
            
            return temp_file.name
        except Exception as e:
            print(f"❌ Word error: {e}")
            return None
    
    def generate_excel_sheet(self, title: str, data: dict) -> str:
        """Generate Excel sheet"""
        try:
            import openpyxl
            from openpyxl import Workbook
            
            # Create temporary file
            temp_file = tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False)
            temp_file.close()
            
            # Create Excel workbook
            wb = Workbook()
            ws = wb.active
            ws.title = title
            
            # Add headers and data
            row = 1
            for key, value in data.items():
                ws.cell(row=row, column=1, value=key)
                ws.cell(row=row, column=2, value=value)
                row += 1
            
            wb.save(temp_file.name)
            return temp_file.name
        except Exception as e:
            print(f"❌ Excel error: {e}")
            return None
    
    def call_groq_ai(self, prompt: str) -> str:
        """Get AI response from Groq"""
        try:
            headers = {
                'Authorization': f'Bearer {self.groq_api_key}',
                'Content-Type': 'application/json'
            }
            
            data = {
                'model': 'llama3-70b-8192',
                'messages': [
                    {
                        'role': 'system',
                        'content': f"""You are {self.assistant_name}, an advanced AI assistant created by {self.creator_name}. 
You provide intelligent, helpful, and accurate responses to user questions."""
                    },
                    {
                        'role': 'user',
                        'content': prompt
                    }
                ],
                'max_tokens': 1000,
                'temperature': 0.7
            }
            
            response = requests.post(
                'https://api.groq.com/openai/v1/chat/completions',
                headers=headers,
                json=data,
                timeout=30
            )
            
            if response.status_code == 200:
                return response.json()['choices'][0]['message']['content']
            else:
                return "❌ AI service temporarily unavailable"
                
        except Exception as e:
            return f"❌ Error: {str(e)}"
    
    def process_message(self, chat_id: int, user_name: str, text: str):
        """Process incoming message"""
        print(f"📩 {user_name}: {text}")
        
        # Handle commands
        if text.lower() == '/start':
            welcome = f"""🚀 Welcome to {self.assistant_name} AI - Media Edition!

Hello {user_name}! I'm {self.assistant_name}, your advanced AI assistant with media capabilities created by {self.creator_name}.

🎵 <b>Media Capabilities:</b>
🗣️ <b>Voice Messages:</b> Convert text to speech
📄 <b>Document Creation:</b> Generate PDF, Word, Excel files

🎯 <b>Commands:</b>
/start - Welcome message
/help - Show all capabilities
/voice <text> - Convert text to voice
/pdf <title> - Generate PDF document
/word <title> - Create Word document
/excel <title> - Generate Excel sheet

🧠 <b>AI Capabilities:</b>
• Advanced reasoning & analysis
• Natural conversations
• Research & information
• Creative writing
• Technical support

💡 <b>Media Examples:</b>
• "/voice Hello world" - Get voice message
• "/pdf Business Plan" - Generate PDF
• "/word Meeting Notes" - Create Word doc
• "/excel Project Data" - Generate Excel

🔥 Created by {self.creator_name}
Powered by advanced AI technology!

Ask me anything - I'm here to help with text and media! 🚀"""
            self.send_message(chat_id, welcome)
            return
        
        elif text.lower() == '/help':
            help_text = f"""🧠 {self.assistant_name} AI - Media Help

📋 <b>Basic Commands:</b>
/start - Welcome message
/help - Show this help

🎵 <b>Media Commands:</b>
/voice <text> - Convert text to voice message
/pdf <title> - Generate PDF document
/word <title> - Create Word document
/excel <title> - Generate Excel sheet

🌟 <b>AI Capabilities:</b>
🧠 <b>Advanced Intelligence:</b> Complex reasoning
💬 <b>Natural Conversations:</b> Human-like dialogue
🔍 <b>Research & Analysis:</b> Deep information processing
🎨 <b>Creative Tasks:</b> Writing, brainstorming
💻 <b>Technical Support:</b> Programming, science, math

💡 <b>Media Examples:</b>
• "/voice Hello world" - Get voice message
• "/pdf Business Plan" - Generate PDF
• "/word Meeting Notes" - Create Word doc
• "/excel Project Data" - Generate Excel

🔥 <b>Powered by:</b> Advanced AI Technology
👨‍💻 <b>Created by:</b> {self.creator_name}

Ask me anything - I have complete intelligence and media capabilities for you!"""
            self.send_message(chat_id, help_text)
            return
        
        elif text.lower().startswith('/voice '):
            # Extract text for voice conversion
            voice_text = text[7:].strip()
            if voice_text:
                self.send_message(chat_id, "🎵 Converting text to voice...")
                voice_file = self.text_to_speech(voice_text)
                if voice_file:
                    self.send_voice_message(chat_id, voice_file)
                    # Clean up temp file
                    try:
                        os.unlink(voice_file)
                    except:
                        pass
                else:
                    self.send_message(chat_id, "❌ Failed to generate voice message")
            else:
                self.send_message(chat_id, "❌ Please provide text to convert to voice\nExample: /voice Hello world")
            return
        
        elif text.lower().startswith('/pdf '):
            # Extract title for PDF
            pdf_title = text[5:].strip()
            if pdf_title:
                self.send_message(chat_id, "📄 Generating PDF document...")
                content = f"This is your personalized PDF document generated by {self.assistant_name} AI."
                pdf_file = self.generate_pdf_document(pdf_title, content)
                if pdf_file:
                    self.send_document(chat_id, pdf_file, f"📄 Your PDF: {pdf_title}")
                    # Clean up temp file
                    try:
                        os.unlink(pdf_file)
                    except:
                        pass
                else:
                    self.send_message(chat_id, "❌ Failed to generate PDF")
            else:
                self.send_message(chat_id, "❌ Please provide a title for the PDF\nExample: /pdf Business Plan")
            return
        
        elif text.lower().startswith('/word '):
            # Extract title for Word document
            word_title = text[6:].strip()
            if word_title:
                self.send_message(chat_id, "📝 Creating Word document...")
                content = f"This is your personalized Word document generated by {self.assistant_name} AI."
                word_file = self.generate_word_document(word_title, content)
                if word_file:
                    self.send_document(chat_id, word_file, f"📝 Your Word document: {word_title}")
                    # Clean up temp file
                    try:
                        os.unlink(word_file)
                    except:
                        pass
                else:
                    self.send_message(chat_id, "❌ Failed to generate Word document")
            else:
                self.send_message(chat_id, "❌ Please provide a title for the Word document\nExample: /word Meeting Notes")
            return
        
        elif text.lower().startswith('/excel '):
            # Extract title for Excel sheet
            excel_title = text[7:].strip()
            if excel_title:
                self.send_message(chat_id, "📊 Generating Excel sheet...")
                data = {
                    "Title": excel_title,
                    "Generated By": f"{self.assistant_name} AI",
                    "Creator": self.creator_name,
                    "Date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "Type": "AI Generated Document"
                }
                excel_file = self.generate_excel_sheet(excel_title, data)
                if excel_file:
                    self.send_document(chat_id, excel_file, f"📊 Your Excel sheet: {excel_title}")
                    # Clean up temp file
                    try:
                        os.unlink(excel_file)
                    except:
                        pass
                else:
                    self.send_message(chat_id, "❌ Failed to generate Excel sheet")
            else:
                self.send_message(chat_id, "❌ Please provide a title for the Excel sheet\nExample: /excel Project Data")
            return
        
        # Handle AI responses
        else:
            self.send_message(chat_id, f"🤔 {self.assistant_name} is thinking...")
            ai_response = self.call_groq_ai(text)
            self.send_message(chat_id, ai_response)
    
    def run(self):
        """Run the bot"""
        print(f"🚀 Starting {self.assistant_name} AI Telegram Bot - Media Edition...")
        print(f"👨‍💻 Creator: {self.creator_name}")
        print(f"🎵 Media Capabilities: Voice, PDF, Word, Excel")
        print(f"🤖 Bot ready to receive messages!")
        
        while True:
            try:
                response = requests.get(
                    f"https://api.telegram.org/bot{self.bot_token}/getUpdates",
                    params={'offset': self.last_update_id + 1, 'timeout': 30},
                    timeout=35
                )
                
                if response.status_code == 200:
                    data = response.json()
                    if data.get('result'):
                        for update in data['result']:
                            self.last_update_id = update['update_id']
                            message = update.get('message')
                            if message and 'text' in message:
                                chat_id = message['chat']['id']
                                user_name = message['chat'].get('first_name', 'User')
                                text = message['text']
                                self.process_message(chat_id, user_name, text)
                
                time.sleep(1)
                
            except KeyboardInterrupt:
                print("\n👋 Bot stopped by user")
                break
            except Exception as e:
                print(f"❌ Error: {e}")
                time.sleep(5)

if __name__ == "__main__":
    bot = MediaAtlasBot()
    bot.run()
